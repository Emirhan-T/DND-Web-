# -*- coding: utf-8 -*-
"""
GrupAdiRapor.docx üretici — DnD Web (Sanal Masaüstü / VTT) Proje Raporu (Türkçe)
python-docx ile oluşturulur.
"""
import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x8a, 0x6d, 0x1a)   # koyu altın
DARK = RGBColor(0x22, 0x22, 0x22)

doc = Document()

# --- Varsayılan stil ---
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_toc(document):
    p = document.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement('w:fldChar'); fld_sep.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = 'İçindekiler tablosunu güncellemek için bu alana sağ tıklayıp "Alanı Güncelle" deyin (veya F9).'
    fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin); run._r.append(instr); run._r.append(fld_sep); run._r.append(t); run._r.append(fld_end)


def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    # arka plan gri benzeri: sınır ekleyelim
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return p


def bullet(text):
    doc.add_paragraph(text, style='List Bullet')


def numbered(text):
    doc.add_paragraph(text, style='List Number')


def kv_table(rows, headers=('Alan', 'Açıklama'), widths=(2.2, 4.3)):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        r = hdr[i].paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        set_cell_bg(hdr[i], '8a6d1a')
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    for i, w in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


# ===================== KAPAK =====================
for _ in range(4):
    doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('DnD Web — Sanal Masaüstü (VTT)')
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = ACCENT

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Yapay Zekâ Entegrasyonlu, Gerçek Zamanlı Çok Oyunculu\nMasaüstü Rol Yapma Oyunu Platformu')
r.font.size = Pt(14)
r.font.color.rgb = DARK

doc.add_paragraph()
sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run('Proje Raporu')
r.bold = True
r.font.size = Pt(16)

for _ in range(6):
    doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Ders: Web Geliştirme\n').bold = True
info.add_run('Grup Adı: [GRUP ADINIZ]\n')
info.add_run('Grup Üyeleri: [Ad Soyad 1], [Ad Soyad 2], [Ad Soyad 3], [Ad Soyad 4]\n')
info.add_run('Teslim Tarihi: ' + datetime.date.today().strftime('%d.%m.%Y') + '\n')
info.add_run('Ankara Üniversitesi')

doc.add_page_break()

# ===================== İÇİNDEKİLER =====================
h = doc.add_heading('İçindekiler', level=1)
add_toc(doc)
doc.add_page_break()

# ===================== 1. GİRİŞ =====================
doc.add_heading('1. Giriş ve Proje Özeti', level=1)
doc.add_paragraph(
    'Bu rapor, “DnD Web” adlı sanal masaüstü (Virtual Tabletop — VTT) uygulamasının '
    'geliştirme sürecini, mimari kararlarını, kullanılan teknolojilerini, veri yapılarını, '
    'yapay zekâ entegrasyonunu ve kullanıcı rehberini kapsamaktadır. Proje, fiziksel olarak '
    'bir araya gelemeyen oyuncuların Dungeons & Dragons (D&D) türü masaüstü rol yapma oyunlarını '
    'çevrim içi, gerçek zamanlı ve etkileşimli biçimde oynayabilmesini amaçlayan tam yığın '
    '(full-stack) bir web uygulamasıdır.'
)
doc.add_paragraph(
    'Uygulamanın merkezinde, bir Oyun Yöneticisi (Game Master — GM) ile birden çok oyuncunun '
    'aynı oyun odasında buluştuğu, gerçek zamanlı senkronize edilen bir oyun paneli '
    '(dashboard) bulunur. Sistem; kullanıcı kimlik doğrulama, ayrıntılı D&D karakter oluşturma, '
    'kampanya/oyun yönetimi, zar atma, sohbet/kayıt günlüğü, savaş (combat) takibi ve bu '
    'raporun özellikle vurguladığı altıgen (hex) tabanlı taktik harita sistemini içerir. '
    'Taktik haritanın arka plan görselleri, OpenAI görüntü üretim API’si ile yapay zekâ '
    'kullanılarak üretilebilmektedir.'
)
doc.add_paragraph(
    'Proje, React tabanlı bir ön yüz (frontend) ile Node.js/Express tabanlı bir arka uç '
    '(backend) ve MongoDB veritabanından oluşur. Gerçek zamanlı iletişim Socket.IO '
    '(WebSocket) ile sağlanır.'
)

# ===================== 2. KAPSAM =====================
doc.add_heading('2. Kapsam', level=1)
doc.add_paragraph(
    'Projenin kapsamı, çok oyunculu bir D&D oturumunun çevrim içi yürütülmesi için gerekli '
    'temel modülleri içerir. Kapsam dâhilindeki ana işlevler şunlardır:'
)
bullet('Kullanıcı yönetimi: kayıt, giriş, JWT tabanlı oturum, “şifremi unuttum” ve e-posta ile şifre sıfırlama.')
bullet('Karakter yönetimi: ayrıntılı D&D karakter sayfası (statlar, yetenekler, büyüler, envanter, biyografi) oluşturma, listeleme, güncelleme, silme.')
bullet('Kampanya/oyun yönetimi: GM tarafından oyun (kampanya) oluşturma, benzersiz oyun kodu ile oyuncuların katılması, kampanya listeleme ve silme.')
bullet('Gerçek zamanlı oyun paneli: GM ve oyuncu panelleri; zar atışı, sohbet/günlük, can (HP) senkronizasyonu, savaş takibi.')
bullet('Altıgen (hex) taktik harita: arazi boyama, hareketli/sabit/toplanabilir nesneler, PC/NPC/canavar token’ları, görünürlük (gizleme) kontrolü, oyuncunun yalnızca kendi token’ını oynatması, lazer işaretçisi.')
bullet('Yapay zekâ entegrasyonu: prompttan taktik harita arka plan görseli üretimi (OpenAI).')
doc.add_paragraph('Kapsam dışı bırakılan / gelecekteki çalışmalar bölümünde ele alınan konular:')
bullet('Otomatik harita VERİSİNİN (arazi/duvar yerleşiminin) yapay zekâ ile üretilmesi (yalnızca arka plan GÖRSELİ üretilir).')
bullet('Sesli/görüntülü iletişim ve mobil yerel (native) uygulama.')

# ===================== 3. AMAÇ =====================
doc.add_heading('3. Amaç', level=1)
doc.add_paragraph('Projenin temel amaçları aşağıda sıralanmıştır:')
numbered('Coğrafi olarak ayrı oyunculara, masaüstü deneyimine yakın, gerçek zamanlı ve etkileşimli bir oyun ortamı sunmak.')
numbered('Oyun Yöneticisine; harita, token ve savaş akışını yöneten merkezî ve yetkili (authoritative) bir kontrol paneli sağlamak.')
numbered('Oyunculara yalnızca görmeleri gereken bilgileri sunarak (görünürlük filtresi) sürpriz/gizlilik mekaniğini korumak.')
numbered('Modern web teknolojileriyle (React, Node, MongoDB, WebSocket) tam yığın bir uygulama geliştirme deneyimi kazanmak.')
numbered('Yapay zekâ servislerini (görüntü üretimi) gerçek bir ürün akışına entegre etmek ve maliyet/performans dengesini gözetmek.')

# ===================== 4. GEREKSİNİMLER =====================
doc.add_heading('4. Gereksinimler', level=1)
doc.add_heading('4.1. Fonksiyonel Gereksinimler', level=2)
kv_table([
    ('FR-1', 'Kullanıcı e-posta ve şifre ile kayıt olabilir ve giriş yapabilir.'),
    ('FR-2', 'Kullanıcı şifresini e-posta yoluyla sıfırlayabilir.'),
    ('FR-3', 'Kullanıcı bir veya daha fazla D&D karakteri oluşturabilir, düzenleyebilir ve silebilir.'),
    ('FR-4', 'GM yeni bir oyun (kampanya) oluşturabilir ve benzersiz bir oyun kodu alır.'),
    ('FR-5', 'Oyuncu, oyun kodu ve bir karakter ile oyuna katılabilir.'),
    ('FR-6', 'GM ve oyuncular aynı oyun odasında gerçek zamanlı senkronize olur.'),
    ('FR-7', 'GM hex haritada arazi boyayabilir, nesne/token yerleştirebilir, öğeleri gizleyebilir.'),
    ('FR-8', 'Oyuncu yalnızca kendi token’ını hareket ettirebilir; gizli öğeleri göremez.'),
    ('FR-9', 'GM, prompt yazarak harita arka planı için yapay zekâ ile görsel üretebilir.'),
    ('FR-10', 'Harita verisi kalıcı olarak saklanır ve sayfa yenilense de geri yüklenir.'),
], headers=('Kod', 'Fonksiyonel Gereksinim'))

doc.add_heading('4.2. Fonksiyonel Olmayan Gereksinimler', level=2)
kv_table([
    ('NFR-1 (Güvenlik)', 'Şifreler bcrypt ile hash’lenir; oturumlar JWT ile doğrulanır; gizli veriler oyunculara hiç gönderilmez.'),
    ('NFR-2 (Performans)', 'Veritabanı yazımları debounce ile (~1,5 sn) toplulaştırılır; gerçek zamanlı olaylar oda bazında yayınlanır.'),
    ('NFR-3 (Kullanılabilirlik)', 'Tek sayfa uygulama (SPA); masaüstü/tablet/mobil duyarlı (responsive) arayüz.'),
    ('NFR-4 (Sürdürülebilirlik)', 'Bileşen tabanlı, ayrık katmanlı mimari; saf yardımcı fonksiyonlar; yeni bağımlılık eklenmeden geliştirme.'),
    ('NFR-5 (Taşınabilirlik)', 'Frontend Vercel/Netlify, backend Render/Railway, veritabanı MongoDB Atlas ile dağıtılabilir.'),
], headers=('Kod', 'Gereksinim'))

# ===================== 5. TEKNOLOJİLER =====================
doc.add_heading('5. Kullanılan Teknolojiler', level=1)
kv_table([
    ('React 19 (CRA)', 'Bileşen tabanlı ön yüz; tek sayfa uygulama (SPA).'),
    ('React Router', 'Sayfa yönlendirme (login, dashboard, karakter sayfaları vb.).'),
    ('SVG', 'Altıgen ızgaranın vektörel çizimi (poligonlar, dönüşüm matrisi).'),
    ('Node.js + Express', 'REST API ve statik dosya sunumu.'),
    ('MongoDB + Mongoose', 'Belge tabanlı veritabanı ve şema/ODM katmanı.'),
    ('Socket.IO', 'WebSocket tabanlı gerçek zamanlı iki yönlü iletişim.'),
    ('JWT (jsonwebtoken)', 'Durumsuz (stateless) kimlik doğrulama.'),
    ('bcryptjs', 'Şifre hash’leme.'),
    ('Nodemailer', 'Şifre sıfırlama e-postaları (Gmail SMTP).'),
    ('OpenAI Images API', 'Yapay zekâ ile arka plan görseli üretimi (gpt-image-1-mini).'),
    ('dotenv', 'Ortam değişkenleri (API anahtarları, gizli bilgiler) yönetimi.'),
], headers=('Teknoloji', 'Kullanım Amacı'))

# ===================== 6. MİMARİ & TASARIM KARARLARI =====================
doc.add_heading('6. Sistem Mimarisi ve Tasarım Kararları', level=1)
doc.add_paragraph(
    'Uygulama üç katmanlı (three-tier) bir mimariye sahiptir: istemci (React), uygulama sunucusu '
    '(Express + Socket.IO) ve veritabanı (MongoDB). Aşağıda öne çıkan tasarım kararları '
    'gerekçeleriyle açıklanmıştır.'
)
doc.add_heading('6.1. Yetkili Sunucu / Pasif İstemci (Authoritative GM)', level=2)
doc.add_paragraph(
    'Oyun durumunun tek doğruluk kaynağı GM’dir. Oyuncular yalnızca istek (örn. “token’ımı şu '
    'hex’e taşı”) gönderir; GM bu isteği doğrular, kendi durumunda uygular ve sonucu odaya geri '
    'yayınlar. Bu model, istemci tarafında hile yapılmasını engeller ve durum tutarlılığını '
    'garanti eder.'
)
doc.add_heading('6.2. Görünürlük Filtresi (Kaynakta Filtreleme)', level=2)
doc.add_paragraph(
    'GM, oyunculara veri yayınlamadan önce toPlayerView() fonksiyonu ile gizli hex’leri, '
    'nesneleri, token’ları ve gmNotes alanını çıkarır. Böylece gizli bilgiler sunucu sınırını '
    'hiç geçmez; tarayıcının ağ sekmesinden bile görülemez. Güvenlik “arayüzde saklama” ile '
    'değil, “kaynakta filtreleme” ile sağlanır.'
)
doc.add_heading('6.3. Paylaşılan Bileşen (Container/Presentational)', level=2)
doc.add_paragraph(
    'Hex harita, GM ve oyuncu panellerinde aynı HexMapBoard bileşeni ile çizilir; role=“gm”|'
    '“player” prop’u neyin gösterileceğini/izin verileceğini belirler. İş mantığı üst bileşenlerde '
    '(container) tutulur, çizim alt bileşende (presentational) yapılır. Bu, kod tekrarını önler.'
)
doc.add_heading('6.4. Altıgen Koordinat Sistemi', level=2)
doc.add_paragraph(
    'Sivri tepeli (pointy-top) altıgenler için eksenel (axial) koordinatlar (q, r) kullanılır. '
    'Piksel dönüşümü: x = size·√3·(q + r/2), y = size·1.5·r. Tıklamadan altıgene geçişte kesirli '
    'koordinatlar küp (cube) yuvarlaması ile en yakın altıgene oturtulur. Izgaranın eğik bir '
    'paralelkenar yerine dikdörtgen olması için ofset koordinat dönüşümü (q = col − floor(r/2)) '
    'uygulanır.'
)
doc.add_heading('6.5. Üç Katmanlı Kalıcılık', level=2)
doc.add_paragraph(
    'Harita durumu üç düzeyde saklanır: (1) localStorage — anında ve yenilemeye dayanıklı; '
    '(2) MongoDB (Game.map) — kalıcı doğruluk kaynağı, debounce ile (~1,5 sn) yazılır; '
    '(3) bellekteki referans — yeni oyuncu katıldığında anında yeniden senkronizasyon için.'
)
doc.add_heading('6.6. Geri Alınabilir / Eklemeli Değişiklikler', level=2)
doc.add_paragraph(
    'Veritabanına eklenen map alanı boş bırakılabilir (nullable); yeni kod ayrı bir klasörde '
    'izole edilir; mevcut Socket.IO olayları değiştirilmeden yalnızca yeni olaylar eklenir. '
    'Böylece sistem her zaman geriye dönük uyumlu kalır.'
)

# ===================== 7. DASHBOARD VERİ YAPISI =====================
doc.add_heading('7. Dashboard (Panel) Veri Yapısı', level=1)
doc.add_paragraph(
    'Bu bölümde paneli besleyen başlıca veri modelleri açıklanır. Veritabanı belgeleri Mongoose '
    'şemaları ile tanımlanır; esnek alanlar Mixed (serbest JSON) tipindedir.'
)

doc.add_heading('7.1. User (Kullanıcı)', level=2)
kv_table([
    ('email', 'String, benzersiz, zorunlu.'),
    ('password', 'String (bcrypt hash), zorunlu.'),
    ('resetPasswordToken', 'String, şifre sıfırlama kodu (geçici).'),
    ('resetPasswordExpire', 'Date, sıfırlama kodunun son geçerlilik anı.'),
    ('createdAt / updatedAt', 'Otomatik zaman damgaları.'),
])

doc.add_heading('7.2. Character (Karakter)', level=2)
kv_table([
    ('userId', 'Karakteri sahibine bağlayan ObjectId.'),
    ('name, level, species, charClass', 'Temel kimlik bilgileri.'),
    ('stats', 'Güç, Çeviklik, Dayanıklılık, Zekâ, Bilgelik, Karizma (Mixed).'),
    ('armorClass, currentHp, maxHp, speed', 'Savaş istatistikleri.'),
    ('weapons, spells, spellSlots', 'Silahlar, büyüler ve büyü slotları (Mixed dizileri).'),
    ('inventory, currency', 'Envanter ve para birimleri (cp/sp/ep/gp/pp).'),
    ('appearance, traits, backstory', 'Biyografi ve rol yapma detayları.'),
])

doc.add_heading('7.3. Game (Oyun / Kampanya)', level=2)
kv_table([
    ('gameCode', 'Benzersiz 8 karakterli oyun kodu (büyük harf).'),
    ('gmId', 'Oyunu kuran GM’in kullanıcı ID’si.'),
    ('name, description', 'Kampanya adı ve açıklaması.'),
    ('isActive', 'Oyunun aktif olup olmadığı.'),
    ('players[]', 'Katılan oyuncular: userId, characterId, playerName, characterName, joinedAt.'),
    ('combatState', 'Savaş durumu: { isActive, round, currentTurnIndex } (Mixed).'),
    ('tokens', 'Eski/uyumlu token dizisi (Mixed).'),
    ('map', 'Hex harita JSON’ı (aşağıda). Yapay zekâ üretimi de buraya yazılabilir.'),
])

doc.add_heading('7.4. Harita JSON Yapısı (Game.map)', level=2)
doc.add_paragraph('Hex harita, src/Components/HexMap/hexUtils.js içindeki buildMapJson() ile aşağıdaki şekilde serileştirilir:')
code_block(
    '{\n'
    '  campaignId, mapName, width, height,\n'
    '  hexes:   [{ q, r, terrain, walkable, movementCost, visibleToPlayers }],\n'
    '  objects: [{ id, type, name, q, r, movable, pickable,\n'
    '              blocksMovement, visibleToPlayers, contains }],\n'
    '  tokens:  [{ id, characterId, ownerId, name, type, q, r,\n'
    '              visibleToPlayers, color, size, hp, maxHp, ac, ... }],\n'
    '  gmNotes, playerDescription, backgroundImageUrl,\n'
    '  hexConfig: { size, visible, opacity }\n'
    '}'
)
doc.add_paragraph(
    'Çalışma zamanında performans için hex’ler “q,r” anahtarlı bir nesne olarak (O(1) erişim), '
    'token’lar ise dizi olarak tutulur. buildMapJson() / applyMapJson() çifti bu iki gösterim '
    'arasında dönüşümü sağlar.'
)
kv_table([
    ('Arazi türleri (terrain)', 'grass, stone, water, mud, road, wall, lava, bridge.'),
    ('Hareketli nesneler', 'barrel, crate, cart, chest.'),
    ('Sabit nesneler', 'tree, wall, rock, door.'),
    ('Toplanabilir nesneler', 'potion, key, scroll, weapon, coin pouch.'),
    ('Token türleri', 'pc, npc, monster.'),
], headers=('Kategori', 'Değerler'))

# ===================== 8. AI ENTEGRASYONU =====================
doc.add_heading('8. Yapay Zekâ (AI) Entegrasyonu', level=1)
doc.add_paragraph(
    'GM, taktik harita için bir sahne tarif eder (örn. “bir nehir ve taş harabe içeren orman '
    'açıklığı”) ve sistem OpenAI görüntü üretim API’si ile yukarıdan görünümlü bir savaş haritası '
    'görseli üretip altıgen ızgaranın arkasına yerleştirir. Akış adımları:'
)
numbered('Ön yüz, prompt ve isteğe bağlı stil/ışık seçeneklerini /api/ai/generate-background uç noktasına gönderir.')
numbered('Backend, sabit bir “şablon” (yukarıdan görünüm, ızgarasız, yazısız, kesintisiz) ile kullanıcının sahnesini birleştirir.')
numbered('OpenAI API’si görüntüyü base64 olarak döndürür; backend bunu PNG dosyasına çözer ve /maps altından statik olarak sunar.')
numbered('Dönen kalıcı URL arka plan olarak ayarlanır, kalıcılaştırılır ve oyunculara senkronize edilir.')
doc.add_paragraph('Prompt şablonu (sabit iskelet + değişken sahne) tasarımı, üretilen tüm görsellerin tutarlı (ızgarasız ve yazısız) olmasını sağlar.')
kv_table([
    ('Model', 'gpt-image-1-mini (varsayılan); gpt-image-1 / dall-e-3 ile değiştirilebilir.'),
    ('Boyut', '1536x1024 (yatay savaş haritası).'),
    ('Yaklaşık maliyet', 'Görsel başına ~0,025 USD (orta kalite); düşük kalitede ~0,005 USD.'),
    ('Sağlayıcı soyutlaması', 'IMAGE_PROVIDER ortam değişkeni ile farklı sağlayıcıya geçiş.'),
    ('Galeri', 'Üretilen son görseller localStorage’da saklanır; tek tıkla yeniden kullanılır.'),
], headers=('Özellik', 'Değer'))
doc.add_paragraph(
    'Önemli tasarım kararı: Görüntünün ham base64 verisi veritabanında saklanmaz veya her '
    'oyuncuya yayınlanmaz; bunun yerine küçük ve kalıcı bir dosya URL’i kullanılır. Bu, ağ ve '
    'depolama maliyetini önemli ölçüde azaltır.'
)

# ===================== 9. KULLANILAN API'LER =====================
doc.add_heading('9. Kullanılan API’ler', level=1)
doc.add_paragraph('Sistem üç tür API kullanır:')
bullet('Dahili REST API’si (Express): /api/auth, /api/characters, /api/games, /api/ai uçları. Ayrıntılar ayrı API dokümantasyonunda (OpenAPI) verilmiştir.')
bullet('Gerçek zamanlı API (Socket.IO): join_room, dice_roll, gm_hexmap_update, player_token_move, laser_point gibi olaylar.')
bullet('Harici API (OpenAI Images): yapay zekâ ile arka plan görseli üretimi.')
doc.add_paragraph('Tüm dahili uç noktaların HTTP method, URL, parametre ve yanıt yapıları, teslim edilen GrupAdiAPI.pdf belgesinde OpenAPI formatında belgelenmiştir.')

# ===================== 10. GERÇEK ZAMANLI & GÜVENLİK =====================
doc.add_heading('10. Gerçek Zamanlı İletişim ve Güvenlik', level=1)
doc.add_paragraph(
    'Her oyun bir Socket.IO “odası” (gameCode) ile temsil edilir. GM bir değişiklik yaptığında '
    'gm_hexmap_update yayınlar; sunucu bunu odadaki oyunculara hexmap_updated olarak iletir. '
    'Oyuncu eylemleri (player_token_move, player_object_pickup) GM’e gider; GM doğrular, uygular '
    've yeniden yayınlar. Lazer işaretçisi (laser_point) kalıcı olmayan, anlık bir olaydır.'
)
doc.add_paragraph('Güvenlik önlemleri:')
bullet('Şifreler bcrypt ile hash’lenir; düz metin saklanmaz.')
bullet('Korumalı uç noktalar JWT (Bearer token) ile doğrulanır (verifyToken ara katmanı).')
bullet('Haritayı yalnızca oyunu kuran GM kaydedebilir (sunucu tarafında gmId kontrolü).')
bullet('Gizli harita verileri toPlayerView() ile sunucuda filtrelenir; oyunculara hiç gönderilmez.')
bullet('Kullanıcı yalnızca kendi karakterlerine erişebilir (userId eşleşmesi).')

# ===================== 11. WEB GELİŞTİRME YAŞAM DÖNGÜSÜ =====================
doc.add_heading('11. Web Geliştirme Yaşam Döngüsü', level=1)
doc.add_paragraph('Proje, klasik yazılım/web geliştirme yaşam döngüsü adımları izlenerek geliştirilmiştir:')
doc.add_heading('11.1. Gereksinim Analizi', level=2)
doc.add_paragraph('Hedef kullanıcılar (GM ve oyuncular) ve temel senaryolar belirlendi; fonksiyonel/fonksiyonel olmayan gereksinimler (Bölüm 4) çıkarıldı.')
doc.add_heading('11.2. Tasarım', level=2)
doc.add_paragraph('Üç katmanlı mimari, veri modelleri (User/Character/Game), REST ve Socket.IO sözleşmeleri, hex koordinat sistemi ve arayüz akışları tasarlandı (Bölüm 6–7).')
doc.add_heading('11.3. Geliştirme', level=2)
doc.add_paragraph('Backend uçları ve şemaları, ardından React bileşenleri ve gerçek zamanlı senkronizasyon kademeli olarak geliştirildi. Hex harita, paylaşılan bir bileşen olarak ayrıştırıldı; AI entegrasyonu sağlayıcı soyutlaması ile eklendi.')
doc.add_heading('11.4. Test', level=2)
doc.add_paragraph('GM ve oyuncu iki ayrı tarayıcı/oturumda açılarak uçtan uca senaryolar denendi: arazi boyama, gizli öğe kontrolü, token hareketi, kalıcılık (yenileme/sunucu yeniden başlatma) ve AI görsel üretimi. Üretim derlemesi (build) uyarı düzeyinde temiz tutuldu.')
doc.add_heading('11.5. Dağıtım (Deployment)', level=2)
doc.add_paragraph('Önerilen dağıtım: Frontend için Vercel/Netlify, backend için Render/Railway, veritabanı için MongoDB Atlas. API anahtarları ortam değişkenleri ile yönetilir; gerçek .env dosyası sürüm kontrolüne eklenmez.')

# ===================== 12. KULLANICI REHBERİ =====================
doc.add_heading('12. Kullanıcı Rehberi', level=1)
doc.add_heading('12.1. Kurulum ve Çalıştırma', level=2)
numbered('Backend: backend klasöründe "npm install" çalıştırın, ardından .env dosyasına MONGO_URI, JWT_SECRET, EMAIL_USER, EMAIL_PASS ve OPENAI_API_KEY değerlerini girin.')
numbered('Backend’i başlatın: backend klasöründe "node server.js" (sunucu 5001 portunda çalışır).')
numbered('Frontend: proje kök klasöründe "npm install" ve ardından "npm start" (uygulama 3000 portunda açılır).')
doc.add_heading('12.2. Hesap ve Karakter', level=2)
numbered('Kayıt olun veya giriş yapın. Şifrenizi unutursanız “Şifremi Unuttum” ile e-posta üzerinden sıfırlayın.')
numbered('Bir veya daha fazla D&D karakteri oluşturun (stat, sınıf, ırk, büyü, envanter vb.).')
doc.add_heading('12.3. Oyun Yöneticisi (GM) Akışı', level=2)
numbered('“Oyun Kur” ile yeni bir kampanya oluşturun; üretilen oyun kodunu oyuncularla paylaşın.')
numbered('Alt sekmelerden HEX HARİTA panelini açın; arazi paletinden hex’leri boyayın, nesne/token yerleştirin.')
numbered('Gizlemek istediğiniz öğeleri “gizli” olarak işaretleyin — oyuncular bunları görmez.')
numbered('İsterseniz “AI Arka Plan” bölümüne bir sahne yazıp “Arka Plan Üret” ile görsel oluşturun, ardından “Hex ızgarayı arka plana oturt” düğmesini kullanın.')
doc.add_heading('12.4. Oyuncu Akışı', level=2)
numbered('Oyun kodu ve bir karakterle oyuna katılın.')
numbered('Yalnızca görünür haritayı görürsünüz; yalnızca kendi token’ınızı sürükleyerek hareket ettirebilirsiniz.')
numbered('Toplanabilir nesnelere tıklayarak alabilir, lazer işaretçisini kullanarak vurgulama yapabilirsiniz.')
doc.add_heading('12.5. AI Arka Plan İpuçları (Sahne Sunumu)', level=2)
bullet('Maliyeti düşürmek için test sırasında OPENAI_IMAGE_QUALITY=low kullanın.')
bullet('Sunumdan önce 2–3 arka planı önceden üretin; bunlar kalıcıdır ve galeriden tek tıkla seçilebilir.')
bullet('gpt-image modeli hesabınızda etkin değilse OPENAI_IMAGE_MODEL=dall-e-3 değerini kullanın.')

# ===================== 13. SONUÇ =====================
doc.add_heading('13. Sonuç ve Gelecek Çalışmalar', level=1)
doc.add_paragraph(
    'DnD Web; kimlik doğrulama, karakter/kampanya yönetimi, gerçek zamanlı çok oyunculu oyun '
    'paneli, altıgen taktik harita ve yapay zekâ destekli görsel üretimini bir araya getiren, '
    'tam yığın bir web uygulamasıdır. Proje; geometri ve SVG ile istemci tarafı çizim, yetkili '
    'bir Express/MongoDB arka ucu, WebSocket senkronizasyonu, sunucu tarafı yetkilendirme ve '
    'maliyet/performans bilinciyle harici bir AI entegrasyonunu örnekler.'
)
doc.add_paragraph('Gelecek çalışmalar:')
bullet('Token’lar için ızgara üzerinde sis-savaşı (fog of war) ve görüş hattı.')
bullet('Yapay zekâ ile harita VERİSİNİN (arazi/duvar) önerilmesi (yarı otomatik).')
bullet('Çok cihazlı oyun için sunulan görsel URL’lerinin LAN/üretim adresine taşınması.')
bullet('Sesli sohbet ve oturum kaydı/oynatma.')

# --- kaydet ---
out = r'c:\Projects\D&D\docs\GrupAdiRapor.docx'
doc.save(out)
print('OK ->', out)
